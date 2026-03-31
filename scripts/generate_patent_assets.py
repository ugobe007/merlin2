"""
Merlin Energy — Patent Filing Asset Generator
Produces:
  Figure_1_System_Architecture.png
  Figure_2_Workflow.png
  Figure_3_Optimization_Engine.png
  Figure_4_UI.png
  Merlin_Patent_Provisional_Main.docx
"""

import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "/Users/robertchristopher/merlin3/patent_assets"
os.makedirs(OUT, exist_ok=True)

PURPLE = "#7C3AED"
INDIGO = "#4F46E5"
DARK   = "#1E1B4B"
TEAL   = "#0D9488"
AMBER  = "#D97706"
SLATE  = "#334155"
LIGHT  = "#F8F7FF"
WHITE  = "#FFFFFF"
GREEN  = "#059669"
RED    = "#DC2626"

# ─────────────────────────────────────────────────────────────────────────────
# FIGURE 1 — SYSTEM ARCHITECTURE
# ─────────────────────────────────────────────────────────────────────────────
def fig1_system_architecture():
    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 12)
    ax.set_facecolor("#0F0A2A")
    fig.patch.set_facecolor("#0F0A2A")
    ax.axis("off")

    def box(x, y, w, h, color, label, sublabel=None, fontsize=9, alpha=0.92):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor=WHITE,
                              linewidth=1.2, alpha=alpha, zorder=3)
        ax.add_patch(rect)
        cy = y + h/2 + (0.15 if sublabel else 0)
        ax.text(x + w/2, cy, label, ha="center", va="center",
                fontsize=fontsize, fontweight="bold", color=WHITE, zorder=4)
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.22, sublabel, ha="center", va="center",
                    fontsize=6.5, color="#CBD5E1", zorder=4)

    def arrow(x1, y1, x2, y2, color=WHITE):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.5), zorder=5)

    def label_zone(x, y, w, h, text, color):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.15",
                              facecolor=color, edgecolor=color,
                              linewidth=0, alpha=0.12, zorder=1)
        ax.add_patch(rect)
        ax.text(x + 0.15, y + h - 0.25, text, fontsize=7.5,
                fontweight="bold", color=color, zorder=2, alpha=0.9)

    # Title
    ax.text(9, 11.55, "MERLIN ENERGY — SYSTEM ARCHITECTURE",
            ha="center", va="center", fontsize=16, fontweight="bold",
            color=WHITE, zorder=6)
    ax.text(9, 11.15, "Patent Application · Figure 1 · Confidential",
            ha="center", va="center", fontsize=9, color="#94A3B8", zorder=6)

    # ── LAYER ZONES ──────────────────────────────────────────────────────────
    label_zone(0.3,  0.3, 17.4, 2.2,  "USER INTERFACE LAYER",          PURPLE)
    label_zone(0.3,  2.7, 17.4, 2.4,  "WIZARD & ORCHESTRATION LAYER",  INDIGO)
    label_zone(0.3,  5.3, 17.4, 2.4,  "QUOTE & PRICING ENGINE LAYER",  TEAL)
    label_zone(0.3,  7.9, 17.4, 2.5,  "INTELLIGENCE & DATA LAYER",     AMBER)
    label_zone(0.3, 10.55, 17.4, 0.55, "EXTERNAL DATA SOURCES",         "#64748B")

    # ── UI LAYER (y 0.5–2.3) ─────────────────────────────────────────────────
    box(0.6,  0.55, 2.6, 1.5, PURPLE,  "Wizard V8",       "5-Step Config UI")
    box(3.5,  0.55, 2.6, 1.5, PURPLE,  "ProQuote Builder","Advanced Config")
    box(6.4,  0.55, 2.6, 1.5, PURPLE,  "Results / Step 5","Tier Presentation")
    box(9.3,  0.55, 2.6, 1.5, "#7E22CE","Deep Dive Panel", "Bank/Pro Model")
    box(12.2, 0.55, 2.6, 1.5, "#6D28D9","Market Intel",    "Dashboard")
    box(15.1, 0.55, 2.4, 1.5, "#5B21B6","Vendor Portal",   "RFQ / Submit")

    # ── WIZARD LAYER (y 2.9–5.0) ─────────────────────────────────────────────
    box(0.6,  3.0, 3.8, 1.6, INDIGO,   "useWizardV8",      "State Orchestrator\nContent-Addressable Cache")
    box(4.7,  3.0, 3.8, 1.6, INDIGO,   "wizardState.ts",   "Reducer / UX_POLICY\nStep Gate Logic")
    box(8.8,  3.0, 3.8, 1.6, "#3730A3","step4Logic.ts",    "buildTiers()\nFour-Layer Sizing")
    box(12.9, 3.0, 4.8, 1.6, "#312E81","MerlinOrchestrator","TrueQuote Auth\nMagicFit → 3 Tiers")

    # ── ENGINE LAYER (y 5.5–7.5) ─────────────────────────────────────────────
    box(0.6,  5.55, 2.8, 1.6, TEAL,    "TrueQuoteV2",      "Quote Engine\nAuth + Validation")
    box(3.7,  5.55, 2.8, 1.6, TEAL,    "MagicFit.ts",      "3-Tier Optimizer\nUpsize Matrices")
    box(6.8,  5.55, 2.8, 1.6, TEAL,    "MarginPolicy",     "3-Layer Pricing\nForbidden UI Methods")
    box(9.9,  5.55, 2.8, 1.6, "#0F766E","benchmarkSources","NREL/EIA Citations\nTrueQuote™ Compliance")
    box(13.0, 5.55, 2.8, 1.6, "#0F766E","unifiedPricing",  "5-Priority Waterfall\nBlend Function")
    box(16.0, 5.55, 1.7, 1.6, "#134E4A","8760 Sim",        "Hourly\nDispatch")

    # ── INTELLIGENCE LAYER (y 8.1–10.2) ──────────────────────────────────────
    box(0.6,  8.15, 2.5, 1.8, AMBER,   "geoIntelligence",  "Solar Grade\nISO/RTO Mapping")
    box(3.4,  8.15, 2.5, 1.8, AMBER,   "marketDataParser", "Dual-Pattern\nRSS Extraction")
    box(6.2,  8.15, 2.5, 1.8, "#B45309","rssAutoFetch",    "Dual-Pipeline\nAI + Alerts")
    box(9.0,  8.15, 2.5, 1.8, "#92400E","priceAlert",      "5-Tier Classification\nBaseline Inference")
    box(11.8, 8.15, 2.5, 1.8, "#78350F","monteCarlo",      "P10/P50/P90\nLHS Sampling")
    box(14.6, 8.15, 2.5, 1.8, "#78350F","specExtraction",  "Hybrid AI+Regex\nRFP Parser")

    # ── EXTERNAL SOURCES (y 10.6–11.0) ──────────────────────────────────────
    sources = ["NREL ATB 2024", "EIA CBECS", "Supabase DB", "GPT-4 AI", "ISO/RTO Feeds", "Vendor API"]
    for i, s in enumerate(sources):
        bx = 0.8 + i * 2.85
        rect = FancyBboxPatch((bx, 10.62), 2.4, 0.38,
                              boxstyle="round,pad=0.05",
                              facecolor="#1E293B", edgecolor="#475569",
                              linewidth=1, alpha=0.95, zorder=3)
        ax.add_patch(rect)
        ax.text(bx + 1.2, 10.81, s, ha="center", va="center",
                fontsize=7, color="#94A3B8", zorder=4)

    # ── ARROWS between layers ─────────────────────────────────────────────────
    # UI → Wizard
    for cx in [1.9, 4.8, 7.7, 10.6]:
        arrow(cx, 2.05, cx, 3.0)
    # Wizard → Engine
    for cx in [2.5, 5.6, 10.7, 15.3]:
        arrow(cx, 4.6, cx, 5.55)
    # Engine → Intelligence
    for cx in [2.0, 5.1, 8.2, 11.3, 14.4]:
        arrow(cx, 5.55, cx, 9.95, color="#94A3B8")
    # Intelligence → External
    for cx in [1.85, 4.65, 7.45, 10.25, 12.85, 15.85]:
        arrow(cx, 8.15, cx, 11.0, color="#475569")

    plt.tight_layout(pad=0.3)
    out = f"{OUT}/Figure_1_System_Architecture.png"
    plt.savefig(out, dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    print(f"✅ {out}")

# ─────────────────────────────────────────────────────────────────────────────
# FIGURE 2 — WIZARD WORKFLOW FLOWCHART
# ─────────────────────────────────────────────────────────────────────────────
def fig2_workflow():
    fig, ax = plt.subplots(figsize=(14, 20))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 20)
    ax.set_facecolor("#0F0A2A")
    fig.patch.set_facecolor("#0F0A2A")
    ax.axis("off")

    ax.text(7, 19.5, "MERLIN ENERGY — WIZARD WORKFLOW",
            ha="center", va="center", fontsize=15, fontweight="bold", color=WHITE)
    ax.text(7, 19.05, "Patent Application · Figure 2 · Confidential",
            ha="center", va="center", fontsize=9, color="#94A3B8")

    def step_box(x, y, w, h, color, title, body=None):
        r = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.12",
                           facecolor=color, edgecolor=WHITE, linewidth=1.4, alpha=0.93, zorder=3)
        ax.add_patch(r)
        ty = y + h/2 + (0.15 if body else 0)
        ax.text(x + w/2, ty, title, ha="center", va="center",
                fontsize=9, fontweight="bold", color=WHITE, zorder=4)
        if body:
            ax.text(x + w/2, y + h/2 - 0.22, body, ha="center", va="center",
                    fontsize=7, color="#CBD5E1", zorder=4)

    def diamond(x, y, w, h, color, label):
        cx, cy = x + w/2, y + h/2
        xs = [cx, cx + w/2, cx, cx - w/2, cx]
        ys = [cy + h/2, cy, cy - h/2, cy, cy + h/2]
        ax.fill(xs, ys, color=color, alpha=0.9, zorder=3)
        ax.plot(xs, ys, color=WHITE, linewidth=1.2, zorder=4)
        ax.text(cx, cy, label, ha="center", va="center",
                fontsize=8, fontweight="bold", color=WHITE, zorder=5)

    def arr(x1, y1, x2, y2, label=None, color=WHITE):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=1.6), zorder=5)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx + 0.15, my, label, fontsize=7, color="#94A3B8", zorder=6)

    def side_note(x, y, text, color=AMBER):
        ax.text(x, y, text, fontsize=6.8, color=color,
                style="italic", va="center", zorder=5)

    # ── MAIN FLOW (center column x≈3.5, w=7) ────────────────────────────────
    # START
    step_box(4.5, 18.3, 5, 0.55, "#1E293B", "START — User opens Merlin Energy")
    arr(7, 18.3, 7, 17.65)

    # STEP 0
    step_box(3.5, 16.85, 7, 0.75, PURPLE, "STEP 0 — Enter Business Name",
             "detectIndustryFromName() · confidence-weighted NER")
    arr(7, 16.85, 7, 16.1)

    # Diamond: confidence ≥ 0.75?
    diamond(4.5, 15.2, 5, 0.85, INDIGO, "confidence ≥ 0.75?")
    arr(7, 16.85, 7, 16.05)
    # Yes → skip step 2
    arr(9.5, 15.625, 11, 15.625, "YES → skip Step 2", color=GREEN)
    step_box(11.1, 15.2, 2.7, 0.85, GREEN, "Auto-populate\nconstraints", None)
    # No → show step 2
    arr(4.5, 15.2, 4.5, 14.5, "NO", color=AMBER)

    # STEP 1
    step_box(3.5, 13.65, 7, 0.75, PURPLE, "STEP 1 — Enter ZIP Code",
             "350ms debounce · 3-tier fallback · parallel API fetch")
    arr(7, 13.65, 7, 12.9)
    side_note(11.0, 14.0, "← Solar grade\n   Utility rates\n   Weather data", TEAL)

    # Intel reveal
    step_box(3.5, 12.1, 7, 0.75, "#1D4ED8", "Progressive Intel Reveal",
             "Promise.allSettled([utility, solar, weather]) → each fills independently")
    arr(7, 12.1, 7, 11.35)

    # STEP 2
    step_box(3.5, 10.55, 7, 0.75, PURPLE, "STEP 2 — Industry Selector",
             "Shown only if confidence < 0.75  ·  12 industries")
    arr(7, 10.55, 7, 9.8)

    # STEP 3
    step_box(3.5, 8.95, 7, 0.8, PURPLE, "STEP 3 — Equipment / Load Profile",
             "Physics answers → calculateUseCasePower()\nNON_POWER_ANSWER_KEYS excluded from recalc")
    arr(7, 8.95, 7, 8.2)
    side_note(11.0, 9.35, "← Background\n   buildTiers()\n   starts HERE", GREEN)

    # STEP 3.5
    step_box(3.5, 7.35, 7, 0.8, "#7C3AED", "STEP 3.5 — Add-Ons",
             "Solar / Generator / EV Charging / Storage Goals\nGrid reliability → auto-enable generator")
    arr(7, 7.35, 7, 6.6)

    # STEP 4
    step_box(3.5, 5.75, 7, 0.8, TEAL, "STEP 4 — System Configuration",
             "MagicFit 3-Tier Optimizer\nUpsize matrix × generation scenario")
    arr(7, 5.75, 7, 5.0)

    # STEP 5
    step_box(3.5, 4.15, 7, 0.8, GREEN, "STEP 5 — Quote Results",
             "Starter / Recommended / Complete tiers\nNPV · IRR · Payback · 25-Year ROI")
    arr(7, 4.15, 7, 3.4)

    # Diamond: export?
    diamond(4.8, 2.55, 4.4, 0.8, INDIGO, "Export?")
    arr(7, 4.15, 7, 3.35)

    # Export branches
    arr(4.8, 2.95, 3.0, 2.95, "Word", color=PURPLE)
    arr(9.2, 2.95, 10.8, 2.95, "Excel / PDF", color=TEAL)
    step_box(1.2, 2.6, 1.7, 0.7, PURPLE, ".docx\n(auth)", None)
    step_box(10.9, 2.6, 2.0, 0.7, TEAL, ".xlsx / .pdf\n(no auth)", None)

    # END
    arr(7, 2.55, 7, 1.8)
    step_box(4.0, 1.15, 6, 0.6, "#1E293B", "TrueQuote™ Compliance Audit")
    arr(7, 1.15, 7, 0.6)
    step_box(4.5, 0.15, 5, 0.42, "#059669", "QUOTE DELIVERED")

    # Back nav note
    side_note(0.3, 6.5, "Back navigation:\ntiersStatus → 'idle'\nif prev < Step 4", RED)

    plt.tight_layout(pad=0.3)
    out = f"{OUT}/Figure_2_Workflow.png"
    plt.savefig(out, dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    print(f"✅ {out}")

# ─────────────────────────────────────────────────────────────────────────────
# FIGURE 3 — OPTIMIZATION ENGINE (MagicFit + Four-Layer + TrueQuote)
# ─────────────────────────────────────────────────────────────────────────────
def fig3_optimization():
    fig, ax = plt.subplots(figsize=(18, 13))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 13)
    ax.set_facecolor("#0F0A2A")
    fig.patch.set_facecolor("#0F0A2A")
    ax.axis("off")

    ax.text(9, 12.6, "MERLIN ENERGY — OPTIMIZATION ENGINE",
            ha="center", fontsize=15, fontweight="bold", color=WHITE)
    ax.text(9, 12.2, "Four-Layer Sizing · MagicFit 3-Tier · TrueQuote Authentication  ·  Patent Application Figure 3",
            ha="center", fontsize=8.5, color="#94A3B8")

    def box(x, y, w, h, fc, title, lines=None, fs=8.5):
        r = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.12",
                           facecolor=fc, edgecolor=WHITE, linewidth=1.3, alpha=0.92, zorder=3)
        ax.add_patch(r)
        n = len(lines) if lines else 0
        ty = y + h - 0.28
        ax.text(x + w/2, ty, title, ha="center", va="top",
                fontsize=fs, fontweight="bold", color=WHITE, zorder=4)
        if lines:
            for i, l in enumerate(lines):
                ax.text(x + w/2, ty - 0.32 - i*0.28, l, ha="center", va="top",
                        fontsize=6.8, color="#CBD5E1", zorder=4)

    def arr(x1, y1, x2, y2, lbl=None, c=WHITE):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=c, lw=1.6), zorder=5)
        if lbl:
            ax.text((x1+x2)/2 + 0.1, (y1+y2)/2, lbl, fontsize=7, color="#94A3B8")

    def zone(x, y, w, h, color, label):
        r = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.2",
                           facecolor=color, edgecolor=color, linewidth=0, alpha=0.1, zorder=1)
        ax.add_patch(r)
        ax.text(x + 0.2, y + h - 0.25, label, fontsize=8, fontweight="bold",
                color=color, alpha=0.85, zorder=2)

    # ─── FOUR LAYER PYRAMID (left half) ─────────────────────────────────────
    zone(0.3, 0.3, 7.8, 11.6, PURPLE, "FOUR-LAYER SIZING ENGINE  (Patent 1)")

    box(0.6, 10.0, 7.2, 1.4, "#5B21B6", "LAYER 1 — LOCATION",
        ["ZIP → peakSunHours → solarGrade (A/A-/B+/B/B-/C+/C/D)",
         "sunFactor = clamp((PSH − 3.0) / 2.5,  0, 1)",
         "B- gate: PSH ≥ 3.5 required for solar eligibility"])

    arr(4.2, 10.0, 4.2, 9.2)

    box(0.6, 8.0, 7.2, 1.5, "#4C1D95", "LAYER 2 — INDUSTRY / FACILITY",
        ["solarPhysicalCapKW per industry type",
         "criticalLoadPct  (IEEE 446 / NEC 517 compliant)",
         "Examples: car_wash=60kW · hotel=225kW · warehouse=819kW"])

    arr(4.2, 8.0, 4.2, 7.2)

    box(0.6, 5.8, 7.2, 1.7, "#3730A3", "LAYER 3 — OPERATIONAL PROFILE",
        ["baseLoadKW · peakLoadKW (measured demand)",
         "EV charger load pre-merged into demand baseline",
         "BESS application intent from questionnaire"])

    arr(4.2, 5.8, 4.2, 4.95)

    box(0.6, 3.5, 7.2, 1.8, INDIGO, "LAYER 4 — GOAL",
        ["save_more: smaller BESS, moderate solar, shorter duration",
         "save_most: balanced NPV/ROI (recommended)",
         "full_power: max BESS + solar + generator, longest duration",
         "Goals adjust sizing weights — data decides, goals guide"])

    arr(4.2, 3.5, 4.2, 2.7)

    box(0.6, 1.0, 7.2, 1.8, TEAL, "SOLAR SIZING FORMULA (Double-Gate)",
        ["solarOptimalKW = physCapKW × sunFactor × goalPenetration",
         "solarFinalKW   = min(solarOptimalKW, physCapKW)",
         "Gate A: solarGrade ≥ B-   (geographic)",
         "Gate B: solarPhysicalCapKW > 0  (facility)   ← BOTH required"])

    # ─── MAGICFIT (center) ──────────────────────────────────────────────────
    zone(8.4, 0.3, 5.0, 11.6, AMBER, "MAGICFIT OPTIMIZER  (Patent 2)")

    arr(7.8, 6.0, 8.6, 6.0, "base calc →")

    box(8.6, 9.5, 4.6, 2.0, "#92400E", "UPSIZE MATRIX",
        ["Generation Combo    Starter / Recommended / Complete  Duration",
         "Full (solar+gen)      1.0×   /   1.0×   /   1.0×     1.0×",
         "Solar Only           1.15×   /  1.25×   /  1.35×     1.5×",
         "Generator Only        1.0×   /   1.1×   /  1.2×      1.0×",
         "UPS Mode (neither)   1.35×   /  1.5×   /  1.65×     2.0×"], fs=7.5)

    arr(10.9, 9.5, 10.9, 8.55)

    box(8.6, 7.1, 4.6, 1.3, AMBER, "SINGLE BASE CALCULATION",
        ["Three tiers derived from ONE financial model call",
         "Ensures internal NPV/IRR consistency across tiers"])

    arr(10.9, 7.1, 10.9, 6.2)

    box(8.6, 4.7, 4.6, 1.8, "#B45309", "THREE TIERS OUTPUT",
        ["STARTER:      conservative, fast payback",
         "RECOMMENDED:  optimal NPV (MagicFit default)",
         "COMPLETE:     max resilience, longest warranty"])

    arr(10.9, 4.7, 10.9, 3.8)

    box(8.6, 2.2, 4.6, 1.7, "#D97706", "CONTENT-ADDRESSABLE CACHE",
        ["Key = JSON.stringify(20+ state params)",
         "Promise deduplication: same key → same Promise",
         "Stale-while-rebuilding: results visible during rebuild",
         "Pre-compute starts at Step 3 → zero wait at Step 5"])

    arr(10.9, 2.2, 10.9, 1.35)
    box(8.6, 0.4, 4.6, 0.9, GREEN, "→ AUTHENTICATED TIERS (TrueQuote™)", fs=8)

    # ─── TRUEQUOTE COMPLIANCE (right) ───────────────────────────────────────
    zone(13.7, 0.3, 4.0, 11.6, TEAL, "TRUEQUOTE™ COMPLIANCE  (Patent 3)")

    arr(13.2, 6.0, 13.9, 6.0, "→")

    box(13.9, 9.6, 3.7, 1.6, "#0F766E", "SOURCE REGISTRY",
        ["Every value linked to typed AUTHORITATIVE_SOURCE:",
         "id · org · pubDate · retrievalDate · confidence",
         "NREL ATB 2024 · NREL StoreFAST · EIA CBECS 2018"])

    arr(15.75, 9.6, 15.75, 8.65)

    box(13.9, 7.2, 3.7, 1.3, TEAL, "PER-VALUE CITATION",
        ["{ value, unit, sourceId, citation,",
         "  confidence, validFrom, validUntil }"])

    arr(15.75, 7.2, 15.75, 6.2)

    box(13.9, 4.8, 3.7, 1.3, "#0D9488", "AUDIT METADATA",
        ["generatedAt · benchmarkVersion",
         "deviations[]: lineItem · benchmarkVal · reason"])

    arr(15.75, 4.8, 15.75, 3.85)

    box(13.9, 2.4, 3.7, 1.35, "#134E4A", "MARGIN POLICY  (Patent 4)",
        ["Layer A: Market / base cost (SSOT)",
         "Layer B: Obtainable reality buffer",
         "Layer C: Sell price (customer-visible)"])

    arr(15.75, 2.4, 15.75, 1.55)

    box(13.9, 0.5, 3.7, 0.95, "#065F46", "_FORBIDDEN_ UI Methods  (Patent 28)",
        ["Runtime exception if UI computes margin"])

    plt.tight_layout(pad=0.3)
    out = f"{OUT}/Figure_3_Optimization_Engine.png"
    plt.savefig(out, dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    print(f"✅ {out}")

# ─────────────────────────────────────────────────────────────────────────────
# FIGURE 4 — UI WIREFRAME (wizard steps)
# ─────────────────────────────────────────────────────────────────────────────
def fig4_ui():
    fig, ax = plt.subplots(figsize=(20, 12))
    ax.set_xlim(0, 20)
    ax.set_ylim(0, 12)
    ax.set_facecolor("#0F0A2A")
    fig.patch.set_facecolor("#0F0A2A")
    ax.axis("off")

    ax.text(10, 11.6, "MERLIN ENERGY — WIZARD USER INTERFACE",
            ha="center", fontsize=15, fontweight="bold", color=WHITE)
    ax.text(10, 11.2, "Patent Application · Figure 4 · Step-by-Step Configuration Flow",
            ha="center", fontsize=9, color="#94A3B8")

    def screen(x, y, w, h, title, step_num, elements, color=PURPLE):
        # Outer frame (phone/tablet bezel)
        bezel = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.12",
                               facecolor="#1E1B4B", edgecolor=color,
                               linewidth=2.5, alpha=0.98, zorder=3)
        ax.add_patch(bezel)

        # Title bar
        bar = FancyBboxPatch((x, y + h - 0.6), w, 0.58,
                             boxstyle="round,pad=0.05",
                             facecolor=color, edgecolor="none",
                             linewidth=0, alpha=0.95, zorder=4)
        ax.add_patch(bar)
        ax.text(x + w/2, y + h - 0.32, title,
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, zorder=5)

        # Step badge
        badge = plt.Circle((x + w - 0.28, y + h - 0.3), 0.22,
                            color=WHITE, zorder=5)
        ax.add_patch(badge)
        ax.text(x + w - 0.28, y + h - 0.3, str(step_num),
                ha="center", va="center", fontsize=8,
                fontweight="bold", color=color, zorder=6)

        # Elements
        ey = y + h - 0.8
        for (etype, text) in elements:
            ey -= 0.38
            if ey < y + 0.08:
                break
            if etype == "label":
                ax.text(x + 0.2, ey, text, fontsize=6.5, color="#94A3B8", zorder=4)
            elif etype == "input":
                r = FancyBboxPatch((x + 0.15, ey - 0.18), w - 0.3, 0.28,
                                   boxstyle="round,pad=0.04",
                                   facecolor="#0F172A", edgecolor="#475569",
                                   linewidth=1, alpha=0.95, zorder=4)
                ax.add_patch(r)
                ax.text(x + 0.3, ey - 0.04, text, fontsize=6.5,
                        color="#64748B", zorder=5)
            elif etype == "card":
                r = FancyBboxPatch((x + 0.15, ey - 0.22), w - 0.3, 0.32,
                                   boxstyle="round,pad=0.04",
                                   facecolor="#1E293B", edgecolor="#334155",
                                   linewidth=1, alpha=0.95, zorder=4)
                ax.add_patch(r)
                ax.text(x + 0.3, ey - 0.06, text, fontsize=6.5,
                        color="#CBD5E1", zorder=5)
            elif etype == "highlight":
                r = FancyBboxPatch((x + 0.15, ey - 0.22), w - 0.3, 0.32,
                                   boxstyle="round,pad=0.04",
                                   facecolor=color, edgecolor="none",
                                   linewidth=0, alpha=0.35, zorder=4)
                ax.add_patch(r)
                ax.text(x + 0.3, ey - 0.06, text, fontsize=6.8,
                        fontweight="bold", color=WHITE, zorder=5)
            elif etype == "btn":
                r = FancyBboxPatch((x + 0.4, ey - 0.18), w - 0.8, 0.28,
                                   boxstyle="round,pad=0.06",
                                   facecolor=color, edgecolor="none",
                                   linewidth=0, alpha=0.95, zorder=4)
                ax.add_patch(r)
                ax.text(x + w/2, ey - 0.04, text, ha="center",
                        fontsize=7, fontweight="bold", color=WHITE, zorder=5)
            ey -= 0.05

    def arr(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color="#7C3AED", lw=2.0), zorder=6)

    sw = 3.4   # screen width
    sh = 10.2  # screen height
    gap = 0.42
    xs = [0.3 + i * (sw + gap) for i in range(5)]
    sy = 0.6

    # STEP 0
    screen(xs[0], sy, sw, sh, "Enter Business Name", 0, [
        ("label",     "What's the business name?"),
        ("input",     "e.g.  Marriott Downtown"),
        ("highlight", "⚡ Detected: Hotel (90%)"),
        ("card",      "Hotel · 225 kW solar cap"),
        ("card",      "Critical load: 45%"),
        ("label",     "Industry auto-populated"),
        ("card",      "Skipping Step 2 →"),
        ("label",     "Patent 11: Confidence gate"),
        ("card",      "getFacilityConstraints()"),
        ("label",     "Atomic constraint populate"),
        ("btn",       "CONFIRM  →"),
    ], PURPLE)

    arr(xs[0] + sw, sy + sh/2, xs[1], sy + sh/2)

    # STEP 1
    screen(xs[1], sy, sw, sh, "ZIP Code & Location Intel", 1, [
        ("label",     "Project ZIP Code"),
        ("input",     "e.g.  90210"),
        ("label",     "350ms debounce · 3-tier fallback"),
        ("highlight", "☀️  Solar Grade: A  (5.8 PSH)"),
        ("card",      "Utility Rate: $0.24/kWh"),
        ("card",      "Peak Demand Rate: $18/kW"),
        ("card",      "ITC Bonus: 10% (California)"),
        ("highlight", "⚡ Grid: Reliable"),
        ("card",      "SREC: Available"),
        ("label",     "Patent 10: Progressive reveal"),
        ("btn",       "ANALYZE LOCATION  →"),
    ], INDIGO)

    arr(xs[1] + sw, sy + sh/2, xs[2], sy + sh/2)

    # STEP 3
    screen(xs[2], sy, sw, sh, "Load Profile & Equipment", 3, [
        ("label",     "Hotel rooms"),
        ("input",     "150  rooms"),
        ("label",     "Sq footage"),
        ("input",     "45,000  sqft"),
        ("highlight", "⚡ Peak Demand:  432 kW"),
        ("card",      "Base Load:   220 kW"),
        ("card",      "EV load pre-merged: +18 kW"),
        ("label",     "Solar roof estimate: 30,000 sqft"),
        ("card",      "Solar Cap: 225 kW"),
        ("label",     "Patent 12: Selective memoize"),
        ("btn",       "CALCULATE LOAD  →"),
    ], "#5B21B6")

    arr(xs[2] + sw, sy + sh/2, xs[3], sy + sh/2)

    # STEP 3.5 / 4
    screen(xs[3], sy, sw, sh, "Add-Ons & Configuration", 4, [
        ("highlight", "✅ Solar PV  (225 kW)"),
        ("card",      "☀️  Grade A · 5.8 PSH"),
        ("card",      "Est. gen:  310 MWh/yr"),
        ("highlight", "✅ Backup Generator"),
        ("card",      "Auto-enabled: unreliable grid"),
        ("card",      "Patent 14: Reducer-level"),
        ("label",     "Goal: Save the Most"),
        ("highlight", "🔋 MagicFit Optimizing..."),
        ("card",      "Upsize matrix: Solar Only"),
        ("card",      "1.15× / 1.25× / 1.35×"),
        ("btn",       "VIEW QUOTES  →"),
    ], TEAL)

    arr(xs[3] + sw, sy + sh/2, xs[4], sy + sh/2)

    # STEP 5
    screen(xs[4], sy, sw, sh, "Quote Results (Step 5)", 5, [
        ("label",     "TrueQuote™ Authenticated"),
        ("highlight", "⭐ RECOMMENDED"),
        ("card",      "BESS:  1.4 MW / 5.6 MWh"),
        ("card",      "Solar: 225 kW"),
        ("card",      "Total: $2,847,000"),
        ("highlight", "💰 Annual Savings: $312K"),
        ("card",      "Payback:  9.1 years"),
        ("card",      "25-Yr NPV: $1.84M"),
        ("card",      "25-Yr ROI:  187%"),
        ("label",     "Export: Word / Excel / PDF"),
        ("btn",       "DOWNLOAD QUOTE  →"),
    ], GREEN)

    # Caption
    ax.text(10, 0.28,
            "Figure 4 — Merlin Energy Wizard Interface: 5-step guided configuration producing TrueQuote™-authenticated BESS system quotes",
            ha="center", fontsize=8, color="#64748B", style="italic")

    plt.tight_layout(pad=0.3)
    out = f"{OUT}/Figure_4_UI.png"
    plt.savefig(out, dpi=180, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    print(f"✅ {out}")

# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT — Provisional Patent Main
# ─────────────────────────────────────────────────────────────────────────────
def make_word_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def style_run(run, size=11, bold=False, color=None, italic=False):
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def h1(text):
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)

    def h2(text):
        p = doc.add_heading(text, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        p.paragraph_format.space_before = Pt(12)

    def h3(text):
        p = doc.add_heading(text, level=3)
        p.paragraph_format.space_before = Pt(8)

    def para(text, indent=0, italic=False):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent   = Inches(indent * 0.25)
        p.paragraph_format.space_after   = Pt(4)
        if italic:
            for r in p.runs:
                r.italic = True
        return p

    def bullet(text, level=0):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
        run = p.add_run(text)
        run.font.size = Pt(10.5)

    def add_fig(path, caption, width=6.0):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(width))
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            cp.paragraph_format.space_after = Pt(14)

    def divider():
        doc.add_paragraph("─" * 80)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    title = doc.add_heading("PROVISIONAL PATENT APPLICATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("MERLIN ENERGY CONFIGURATION SYSTEM")
    style_run(r, size=16, bold=True, color=(0x7C, 0x3A, 0xED))

    doc.add_paragraph()
    meta = [
        ("Applicant:",       "Robert Christopher"),
        ("Product:",         "Merlin Energy — Commercial BESS Quoting & Configuration Platform"),
        ("Live URL:",        "https://merlinenergy.net"),
        ("Filing Date:",     "March 30, 2026"),
        ("Document Type:",   "Provisional Patent Application (35 U.S.C. § 111(b))"),
        ("Classification:",  "CONFIDENTIAL — ATTORNEY-CLIENT PRIVILEGED"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    note = doc.add_paragraph(
        "⚠️  This document is attorney-client privileged and contains trade secret and "
        "patent-sensitive information. Do not distribute outside legal counsel.")
    note.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    note.runs[0].bold = True

    doc.add_page_break()

    # ── TITLE OF INVENTION ────────────────────────────────────────────────────
    h1("TITLE OF INVENTION")
    para("AUTOMATED COMMERCIAL ENERGY STORAGE SYSTEM CONFIGURATION AND QUOTING PLATFORM "
         "WITH FOUR-LAYER PHYSICS SIZING, PROACTIVE BACKGROUND SYNTHESIS, AND "
         "SOURCE-TRACEABLE PRICING COMPLIANCE")

    # ── CROSS REFERENCES ──────────────────────────────────────────────────────
    h1("CROSS-REFERENCE TO RELATED APPLICATIONS")
    para("This is an original provisional patent application. No prior provisional applications "
         "have been filed for this invention. Applicant intends to file a non-provisional "
         "application claiming priority to this provisional within twelve (12) months of "
         "this filing date pursuant to 35 U.S.C. § 119(e).")

    # ── FIELD ─────────────────────────────────────────────────────────────────
    h1("FIELD OF THE INVENTION")
    para("The present invention relates to computer-implemented systems for configuring and "
         "pricing commercial battery energy storage systems (BESS), and more particularly "
         "to automated multi-step configuration platforms that combine geographic intelligence, "
         "physics-based facility sizing, machine-learning-assisted industry classification, "
         "proactive background computation, and source-traceable financial modeling to "
         "generate bankable commercial energy system quotes.")

    # ── BACKGROUND ────────────────────────────────────────────────────────────
    h1("BACKGROUND OF THE INVENTION")
    para("Commercial battery energy storage systems represent a rapidly growing segment of "
         "the energy market. However, properly sizing a BESS for a commercial facility "
         "requires expertise across multiple domains: electrical engineering (demand "
         "calculations), solar irradiance analysis, utility rate structures, financial "
         "modeling, and equipment pricing intelligence.")

    para("Existing solutions suffer from significant deficiencies:")
    for d in [
        "Manual quoting processes requiring 2–4 weeks and specialized engineering staff",
        "Single-point sizing that produces one configuration rather than optimized alternatives",
        "Static pricing databases that do not incorporate real-time market intelligence",
        "No source attribution — customers cannot verify the basis for quoted numbers",
        "Separate tools for sizing, pricing, and financial analysis requiring manual integration",
        "No probabilistic risk modeling for financial performance guarantees",
    ]:
        bullet(d)

    para("\nWhat is needed is a unified automated platform that simultaneously addresses "
         "all of these deficiencies while producing bankable, source-traceable quotes "
         "suitable for financing and investment decisions.")

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    h1("SUMMARY OF THE INVENTION")
    para("The present invention provides a computer-implemented system for automatically "
         "generating bankable commercial energy storage system configurations, comprising:")

    claims_summary = [
        "A location-aware geographic intelligence layer that resolves a user-supplied "
         "location identifier through a multi-tier fallback chain while simultaneously "
         "initiating independent parallel fetches for utility rate data, solar irradiance "
         "data, and weather data;",
        "A probabilistic business classification engine that maps a business name to a "
         "facility-type category using confidence-weighted keyword matching, eliminates "
         "intermediate configuration steps when classification confidence exceeds a threshold, "
         "and atomically populates physics constraints in the same transaction;",
        "A four-layer energy system sizing engine that combines geographic, industry, "
         "operational, and goal parameters into three distinct optimized configurations;",
        "A proactive synthesis engine that begins computing configurations before the user "
         "requests them using a content-addressable cache;",
        "A source-traceable pricing engine where every computed value is linked to a "
         "documented external authority with full audit metadata.",
    ]
    for i, c in enumerate(claims_summary, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        r = p.add_run(f"({i})  ")
        r.bold = True
        p.add_run(c).font.size = Pt(10.5)

    # ── FIGURES ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1("BRIEF DESCRIPTION OF THE DRAWINGS")

    figs = [
        ("Figure 1", "Figure_1_System_Architecture.png",
         "System Architecture — Seven-layer architecture showing the User Interface Layer, "
         "Wizard & Orchestration Layer, Quote & Pricing Engine Layer, Intelligence & Data "
         "Layer, and External Data Sources with data flow arrows between components."),
        ("Figure 2", "Figure_2_Workflow.png",
         "Wizard Workflow Flowchart — End-to-end user flow from business name entry through "
         "ZIP code resolution, industry classification, load profiling, add-on selection, "
         "MagicFit optimization, and TrueQuote™-authenticated quote delivery with export options."),
        ("Figure 3", "Figure_3_Optimization_Engine.png",
         "Optimization Engine — Three-panel diagram showing the Four-Layer Sizing Engine "
         "(Layers 1–4 with solar double-gate formula), MagicFit Three-Tier Optimizer "
         "(upsize matrix and content-addressable cache), and TrueQuote™ Compliance stack "
         "(source registry, audit metadata, margin policy, and Forbidden UI methods)."),
        ("Figure 4", "Figure_4_UI.png",
         "User Interface — Five-screen wireframe depicting Steps 0 through 5 of the "
         "Merlin Energy wizard, showing business detection, location intel reveal, "
         "load profile entry, add-on configuration, and final quote results with "
         "export options."),
    ]

    for fig_name, fig_file, fig_desc in figs:
        p = doc.add_paragraph()
        r = p.add_run(fig_name + ":  ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        p.add_run(fig_desc).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()
    h1("DRAWINGS")
    for fig_name, fig_file, fig_desc in figs:
        h3(fig_name)
        add_fig(f"{OUT}/{fig_file}", f"{fig_name} — {fig_desc[:80]}...", width=6.2)

    # ── DETAILED DESCRIPTION ─────────────────────────────────────────────────
    doc.add_page_break()
    h1("DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS")

    patents = [
        ("Patent 1 — Four-Layer Energy System Sizing Engine",
         "The sizing engine operates as a four-layer information pyramid where each "
         "layer is collected sequentially and each constrains the next. Layer 1 derives "
         "a continuous solar viability coefficient: sunFactor = clamp((peakSunHours − 3.0) "
         "/ 2.5, 0, 1). A letter-grade system (A through D) maps PSH values to solar "
         "eligibility; B− (PSH ≥ 3.5) is the minimum viable grade constituting a hard "
         "exclusion gate. Layer 2 applies industry-specific physical constraints including "
         "maximum solar hosting capacity and critical load percentage per IEEE 446-1995. "
         "Layer 3 incorporates measured operational demand with EV charging pre-merged "
         "into the baseline. Layer 4 applies goal-weighted penetration adjustments.",
         ["src/wizard/v8/step4Logic.ts",
          "src/wizard/v8/wizardState.ts (isSolarFeasible)",
          "src/services/useCasePowerCalculations.ts"]),

        ("Patent 2 — MagicFit: Generative Three-Tier Optimizer",
         "MagicFit generates three distinct optimized energy system configurations from "
         "a single base calculation. A four-scenario upsize matrix dynamically adjusts "
         "BESS capacity based on whether solar and/or generator are included: Full "
         "generation (1.0×/1.0×/1.0×), Solar Only (1.15×/1.25×/1.35× with 1.5× duration), "
         "Generator Only (1.0×/1.1×/1.2×), and UPS Mode — when neither source is selected "
         "— (1.35×/1.5×/1.65× with 2.0× duration). All three tiers derive from a single "
         "financial model call ensuring internal NPV/IRR consistency.",
         ["src/services/MagicFit.ts",
          "src/services/TrueQuoteEngineV2.ts",
          "src/services/MerlinOrchestrator.ts"]),

        ("Patent 3 — Source-Traceable Quote Generation (TrueQuote™)",
         "Every computed value in a customer-facing quote is linked to a documented, "
         "version-controlled external authority. A typed AUTHORITATIVE_SOURCES registry "
         "contains id, name, organization, type, publicationDate, retrievalDate, vintage, "
         "lastVerified, url, and notes for each source. Per-value citations return "
         "{ value, unit, sourceId, citation, confidence, validFrom, validUntil }. "
         "QuoteAuditMetadata includes deviations[] logging any value that differs from "
         "the national benchmark with lineItem, benchmarkValue, appliedValue, and reason.",
         ["src/services/benchmarkSources.ts",
          "src/services/marginPolicyEngine.ts",
          "src/services/unifiedQuoteCalculator.ts"]),

        ("Patent 9 — Proactive Background Synthesis Engine",
         "The system proactively begins computing final quote results at Step 3, two "
         "full steps before the results page. A content-addressable cache key "
         "(createTierBuildKey) is a deterministic JSON fingerprint derived from 20+ "
         "state parameters including location, industry, load measurements, add-on "
         "selections, and geographic intelligence values. If the fingerprint matches an "
         "in-flight Promise, that Promise is returned without duplication. Stale-while-"
         "rebuilding ensures existing results remain visible during background rebuilds.",
         ["src/wizard/v8/useWizardV8.ts (createTierBuildKey)",
          "src/wizard/v8/step4Logic.ts (buildTiers)"]),

        ("Patent 11 — Confidence-Gated Wizard Step Elimination",
         "When business name classification confidence exceeds 0.75, the system "
         "simultaneously eliminates Step 2 from navigation AND populates all downstream "
         "physics constraints (solarPhysicalCapKW, criticalLoadPct) in a single atomic "
         "two-phase dispatch. This prevents the null-state window that would otherwise "
         "exist between step elimination and constraint availability — a window that "
         "would cause solar to be excluded from all subsequent calculations. The client-"
         "side classifier uses confidence-weighted keyword matching with named entity "
         "recognition for 30+ industry categories and major brand names.",
         ["src/wizard/v8/useWizardV8.ts (detectIndustryFromName)",
          "src/wizard/v8/wizardState.ts (CONFIRM_BUSINESS reducer)"]),

        ("Patent 28 — Forbidden UI Computation Enforcement",
         "The pricing engine attaches _FORBIDDEN_computeMarginInUI() and "
         "_FORBIDDEN_computeNetCostInUI() methods to every MarginRenderEnvelope result "
         "object. These methods throw runtime exceptions if called, making it "
         "structurally impossible for UI code to compute margin independently. This "
         "architectural enforcement guarantees that every sell price reaching the "
         "customer interface was produced by the marginPolicyEngine with full audit trail. "
         "The _FORBIDDEN_ naming convention also enables static analysis detection.",
         ["src/services/marginPolicyEngine.ts (MarginRenderEnvelope)",
          "src/services/TrueQuoteEngineV2.ts"]),
    ]

    for title, desc, sources in patents:
        h2(title)
        para(desc)
        p = doc.add_paragraph()
        r = p.add_run("Key Source Files:  ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        for s in sources:
            bullet(s, level=1)
        doc.add_paragraph()

    # ── CLAIMS ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1("CLAIMS")
    para("What is claimed is:", italic=True)
    doc.add_paragraph()

    claims = [
        ("Independent Claim 1",
         "A computer-implemented system for automatically generating commercial energy "
         "storage system configurations, comprising: a processor; a non-transitory "
         "computer-readable medium storing instructions that, when executed, cause the "
         "processor to: receive a postal code from a user; resolve the postal code "
         "through a multi-tier fallback chain comprising an external geocoding service, "
         "a local utility rate database, and an identity fallback; simultaneously "
         "initiate independent parallel data fetches for solar irradiance data, utility "
         "rate data, and weather data; and populate a geographic intelligence profile "
         "independently from each fetch result as it completes."),
        ("Independent Claim 2",
         "A computer-implemented method for sizing a commercial battery energy storage "
         "system, comprising: deriving a solar viability coefficient from a peak sun "
         "hours value using a continuous transfer function; applying a letter-grade "
         "classification wherein a minimum grade threshold constitutes a hardware "
         "exclusion gate preventing solar inclusion in system configurations; applying "
         "industry-specific physical capacity constraints; incorporating measured "
         "operational demand data; and applying goal-weighted sizing adjustments."),
        ("Independent Claim 3",
         "A computer-implemented system for generating optimized energy storage "
         "configurations, comprising: a base configuration calculator; a generation "
         "scenario detector that identifies which combination of solar generation and "
         "backup generation sources are selected; an upsize matrix that applies "
         "scenario-specific capacity multipliers to battery storage size and duration, "
         "wherein a configuration with neither solar nor backup generation receives a "
         "maximum upsize multiplier; and a tier generator that derives three distinct "
         "configurations from a single base financial model computation."),
        ("Dependent Claim 4",
         "The system of claim 3, wherein the upsize matrix applies a duration multiplier "
         "of 2.0 to configurations with neither solar nor backup generation, compensating "
         "for the absence of on-site generation through extended storage duration."),
        ("Independent Claim 5",
         "A computer-implemented system for enforcing pricing audit integrity, comprising: "
         "a pricing engine that computes sell prices with full audit metadata; and a result "
         "object returned to a user interface layer comprising embedded methods that "
         "throw runtime exceptions when invoked, wherein the embedded methods correspond "
         "to margin and net cost computation operations, thereby making independent "
         "user-interface-layer pricing computation structurally impossible."),
        ("Independent Claim 6",
         "A computer-implemented method for proactive energy system configuration "
         "synthesis, comprising: monitoring user configuration state during a multi-step "
         "wizard process; computing a deterministic cache key from a plurality of state "
         "parameters; initiating a background computation of final system configurations "
         "responsive to detecting a trigger step, wherein the trigger step precedes the "
         "results step by at least two steps; returning an in-flight computation promise "
         "when the cache key matches an existing computation; and displaying previously "
         "computed results while a background recomputation proceeds."),
        ("Independent Claim 7",
         "A computer-implemented method for business-type-gated wizard navigation, "
         "comprising: applying a confidence-weighted keyword classifier to a business "
         "name input; comparing a resulting confidence score to a threshold; in response "
         "to the confidence score meeting the threshold, in a single atomic transaction: "
         "modifying a navigation state to eliminate an intermediate configuration step "
         "from the user's path, and populating physics constraint parameters required "
         "by subsequent steps."),
    ]

    for i, (claim_title, claim_text) in enumerate(claims, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}.  [{claim_title}]  ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        p.add_run(claim_text).font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(8)

    # ── ABSTRACT ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1("ABSTRACT")
    para(
        "A computer-implemented system and method for automatically generating bankable "
        "commercial energy storage system (BESS) configurations. The system comprises a "
        "location-aware geographic intelligence layer that resolves postal codes through "
        "a multi-tier fallback chain while initiating parallel data fetches; a confidence-"
        "weighted business classification engine that eliminates configuration steps and "
        "atomically populates physics constraints when classification confidence exceeds "
        "a threshold; a four-layer sizing engine combining geographic, industry, operational, "
        "and goal parameters into three optimized configurations via a generation-scenario-"
        "specific upsize matrix; a proactive background synthesis engine using a "
        "content-addressable cache keyed on configuration state to pre-compute results "
        "before the user reaches the results step; a source-traceable pricing engine "
        "linking every computed value to a documented external authority; and a "
        "pricing integrity enforcement mechanism that embeds exception-throwing methods "
        "on result objects to make user-interface-layer margin computation structurally "
        "impossible. The system produces TrueQuote™-authenticated quotes suitable for "
        "commercial energy financing and investment decisions."
    )

    # ── INVENTOR DECLARATION ──────────────────────────────────────────────────
    doc.add_page_break()
    h1("INVENTOR DECLARATION")
    para("I hereby declare that:")
    for d in [
        "I am the original inventor of the subject matter claimed in this application.",
        "I have reviewed and understand the contents of this application.",
        "I acknowledge my duty to disclose to the USPTO all information known to me "
         "to be material to patentability.",
        "I believe the invention described herein has not been publicly disclosed "
         "more than one year prior to this filing date.",
    ]:
        bullet(d)
    doc.add_paragraph()
    para("Inventor:  Robert Christopher")
    para("Date:  March 30, 2026")
    para("Signature:  ___________________________")
    doc.add_paragraph()
    divider()
    p = doc.add_paragraph()
    r = p.add_run(
        f"Document generated: {datetime.datetime.now().strftime('%B %d, %Y')}  ·  "
        "Merlin Energy Platform  ·  CONFIDENTIAL — ATTORNEY-CLIENT PRIVILEGED"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = f"{OUT}/Merlin_Patent_Provisional_Main.docx"
    doc.save(out)
    print(f"✅ {out}")


# ── RUN ALL ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating patent filing assets…")
    fig1_system_architecture()
    fig2_workflow()
    fig3_optimization()
    fig4_ui()
    make_word_doc()
    print(f"\n🎉  All files written to:  {OUT}/")
    import subprocess
    subprocess.run(["ls", "-lh", OUT])
